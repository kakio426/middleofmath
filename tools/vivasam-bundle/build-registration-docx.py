#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "artifacts" / "vivasam" / "vivasam-registration-input-01-03.docx"
PRODUCTION_EVIDENCE = ROOT / "artifacts" / "vivasam" / "eduitit-production-publication.json"
BLUE = RGBColor(46, 116, 181)
INK = RGBColor(31, 45, 61)
MUTED = RGBColor(92, 103, 112)

LESSONS = [
    {
        "number": "01",
        "lesson_id": "g3s2-pictograph-legend",
        "title": "그림 하나에 숨은 수",
        "target": "초등 3학년 2학기 · 6. 그림그래프",
        "intent": (
            "그림 수만 세고 답을 말하는 데서 그치지 않고, 범례가 무엇을 뜻하는지 먼저 읽은 뒤 "
            "그림 수와 단위 수를 연결해 실제 수량을 구하도록 구성했습니다. 서로 다른 답을 비교하며 "
            "어디에서 생각이 달라졌는지 짝에게 설명하고, 활동지에는 식·그림·짧은 문장 가운데 한 가지로 "
            "근거를 남기게 해 범례를 제대로 적용하는지 살펴보고자 했습니다."
        ),
        "url": "https://eduitit.site/edu-materials/2e02ca70-e3ee-4734-9741-bf4a3a5cb7ae/",
        "image": ROOT / "artifacts" / "vivasam" / "g3s2-pictograph-legend" / "support" / "representative-image.png",
    },
    {
        "number": "02",
        "lesson_id": "g3s1-multiplication-groups-model",
        "title": "같은 묶음은 곱셈으로",
        "target": "초등 3학년 1학기 · 1. 곱셈",
        "intent": (
            "같은 수가 여러 묶음으로 놓인 상황을 덧셈과 곱셈으로 자연스럽게 연결하는 데 초점을 두었습니다. "
            "학생이 한 묶음의 수와 묶음 수를 구분해 그림과 식으로 나타내고, 두 풀이가 어디에서 달라졌는지 "
            "서로 설명하게 했습니다. 마지막에는 새로운 묶음 상황을 활동지에서 스스로 표현하게 해 곱셈식을 "
            "계산 결과가 아니라 상황을 나타내는 방법으로 이해했는지 확인하고자 했습니다."
        ),
        "url": "https://eduitit.site/edu-materials/c058e808-6713-42fa-8905-24dc0515776e/",
        "image": ROOT / "artifacts" / "vivasam" / "g3s1-multiplication-groups-model" / "support" / "representative-image.png",
    },
    {
        "number": "03",
        "lesson_id": "g3s1-multiplication-array-transfer",
        "title": "줄과 칸으로 전체 수 찾기",
        "target": "초등 3학년 1학기 · 1. 곱셈",
        "intent": (
            "배열을 보자마자 두 수를 더하지 않고, 한 줄에 몇 개씩 몇 줄이 있는지 차례로 짚어 보게 했습니다. "
            "처음에는 같은 그림을 두고 30과 11이라는 서로 다른 답을 비교하며 각 수가 무엇을 뜻하는지 말하게 하고, "
            "이후에는 연필 봉지와 바둑돌처럼 배치가 달라져도 한 묶음의 수와 묶음 수를 찾아 곱셈식으로 옮기게 했습니다. "
            "마지막에는 잘못된 풀이에서 처음 어긋난 지점을 직접 고쳐 쓰게 해, 계산보다 상황을 읽는 힘이 남도록 구성했습니다."
        ),
        "url": "",
        "image": ROOT / "artifacts" / "vivasam" / "g3s1-multiplication-array-transfer" / "support" / "representative-image.png",
    },
]


def production_urls():
    if not PRODUCTION_EVIDENCE.exists():
        return {}
    import json

    payload = json.loads(PRODUCTION_EVIDENCE.read_text(encoding="utf-8"))
    return {
        record["lessonId"]: record["publicUrl"]
        for record in payload.get("records", [])
        if record.get("lessonId") and record.get("publicUrl")
    }


def set_font(run, size=11, bold=False, color=INK):
    run.font.name = "Pretendard"
    fonts = run._element.get_or_add_rPr().rFonts
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), "Pretendard")
    fonts.set(qn("w:hint"), "eastAsia")
    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "ko-KR")
    language.set(qn("w:eastAsia"), "ko-KR")
    run._element.get_or_add_rPr().append(language)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_paragraph(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_hyperlink(paragraph, text, url):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(properties)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_label(doc, label):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=10, after=4, line=1.0)
    set_font(paragraph.add_run(label), size=11, bold=True, color=BLUE)
    return paragraph


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_font(run, size=9, color=MUTED)


def build():
    urls = production_urls()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Pretendard"
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        normal._element.rPr.rFonts.set(qn(key), "Pretendard")
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(header, after=0, line=1.0)
    set_font(header.add_run("수업 자료 등록용 입력 내용"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(footer, after=0, line=1.0)
    add_page_number(footer)

    for index, lesson in enumerate(LESSONS):
        if index:
            doc.add_page_break()

        kicker = doc.add_paragraph()
        set_paragraph(kicker, after=4, line=1.0)
        set_font(kicker.add_run(f"자료 {lesson['number']}"), size=10, bold=True, color=BLUE)

        title = doc.add_paragraph()
        set_paragraph(title, after=4, line=1.05)
        set_font(title.add_run(lesson["title"]), size=23, bold=True, color=INK)

        meta = doc.add_paragraph()
        set_paragraph(meta, after=12, line=1.0)
        set_font(meta.add_run(lesson["target"]), size=10.5, color=MUTED)

        add_label(doc, "교과목")
        subject = doc.add_paragraph()
        set_paragraph(subject, after=8)
        set_font(subject.add_run("수학"), size=11)

        add_label(doc, "수업 설계 의도")
        intent = doc.add_paragraph()
        set_paragraph(intent, after=8, line=1.35)
        set_font(intent.add_run(lesson["intent"]), size=10.5)

        add_label(doc, "URL 등록")
        url_paragraph = doc.add_paragraph()
        set_paragraph(url_paragraph, after=8, line=1.0)
        url = urls.get(lesson["lesson_id"], lesson["url"])
        if not url:
            raise RuntimeError(f"운영 공개 URL이 없습니다: {lesson['lesson_id']}")
        add_hyperlink(url_paragraph, url, url)

        add_label(doc, "이미지 등록")
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(picture, after=4, line=1.0)
        picture.add_run().add_picture(str(lesson["image"]), width=Inches(5.05))
        image_path = doc.add_paragraph()
        image_path.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(image_path, after=0, line=1.0)
        set_font(image_path.add_run(str(lesson["image"].relative_to(ROOT))), size=8.5, color=MUTED)

    doc.core_properties.title = "수업 자료 등록용 입력 내용 01-03"
    doc.core_properties.subject = "교과목, 수업 설계 의도, 공개 URL, 대표 이미지"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

#!/usr/bin/env python3
"""Build the single copy-and-paste document used for Vivasam registration."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
REGISTER_PATH = REPO_ROOT / "artifacts" / "vivasam" / "vivasam-submission-register.json"
OUTPUT_PATH = REPO_ROOT / "artifacts" / "vivasam" / "비바샘-등록용-문구-30개.docx"

PURPLE = "6D46C7"
LIGHT_PURPLE = "F4F0FF"
DARK = RGBColor(31, 31, 35)
MUTED = RGBColor(91, 91, 102)
DOCUMENT_FONT = "Arial Unicode MS"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=140, start=180, bottom=140, end=180) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size: float, bold: bool = False, color: RGBColor = DARK) -> None:
    run.font.name = DOCUMENT_FONT
    run_properties = run._element.get_or_add_rPr()
    run_properties.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
    run_properties.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
    run_properties.rFonts.set(qn("w:eastAsia"), DOCUMENT_FONT)
    run_properties.rFonts.set(qn("w:cs"), DOCUMENT_FONT)
    run_properties.rFonts.set(qn("w:hint"), "eastAsia")
    language = run_properties.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        run_properties.append(language)
    language.set(qn("w:val"), "ko-KR")
    language.set(qn("w:eastAsia"), "ko-KR")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_label_value(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, 10.5, bold=True, color=RGBColor(78, 53, 140))
    value_run = paragraph.add_run(value)
    set_run_font(value_run, 10.5, color=DARK)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), PURPLE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page(document: Document, record: dict, index: int, total: int) -> None:
    if index:
        document.add_page_break()

    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(3)
    eyebrow_run = eyebrow.add_run(f"VIVASAM REGISTRATION · {record['sequence']:02d}/{total:02d}")
    set_run_font(eyebrow_run, 9, bold=True, color=RGBColor(109, 70, 199))

    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(8)
    heading_run = heading.add_run(record["title"])
    set_run_font(heading_run, 20, bold=True, color=DARK)

    add_label_value(document, "교과목", record["subject"])
    add_label_value(document, "대상·단원", f"{record['grade']} · {record['unit']}")

    intent_label = document.add_paragraph()
    intent_label.paragraph_format.space_before = Pt(8)
    intent_label.paragraph_format.space_after = Pt(4)
    set_run_font(intent_label.add_run("수업 설계 의도"), 11, bold=True, color=RGBColor(78, 53, 140))

    intent_table = document.add_table(rows=1, cols=1)
    intent_table.autofit = False
    intent_table.columns[0].width = Cm(17.2)
    intent_cell = intent_table.cell(0, 0)
    intent_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(intent_cell, LIGHT_PURPLE)
    set_cell_margins(intent_cell, top=220, start=240, bottom=220, end=240)
    intent_paragraph = intent_cell.paragraphs[0]
    intent_paragraph.paragraph_format.line_spacing = 1.35
    intent_paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(intent_paragraph.add_run(record["teachingIntent"]), 11, color=DARK)

    url_label = document.add_paragraph()
    url_label.paragraph_format.space_before = Pt(10)
    url_label.paragraph_format.space_after = Pt(3)
    set_run_font(url_label.add_run("URL 등록"), 11, bold=True, color=RGBColor(78, 53, 140))
    url_paragraph = document.add_paragraph()
    url_paragraph.paragraph_format.space_after = Pt(8)
    add_hyperlink(url_paragraph, record["publicUrl"], record["publicUrl"])

    image_label = document.add_paragraph()
    image_label.paragraph_format.space_after = Pt(5)
    set_run_font(image_label.add_run("이미지 등록"), 11, bold=True, color=RGBColor(78, 53, 140))

    image_reference = record["representativeImagePath"]
    image_path = Path(image_reference.replace("middleofmath:", str(REPO_ROOT) + "/"))
    image_table = document.add_table(rows=1, cols=2)
    image_table.autofit = False
    image_table.columns[0].width = Cm(6.0)
    image_table.columns[1].width = Cm(11.2)
    left, right = image_table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    picture_paragraph = left.paragraphs[0]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Cm(5.4))

    filename_paragraph = right.paragraphs[0]
    filename_paragraph.paragraph_format.space_after = Pt(5)
    set_run_font(filename_paragraph.add_run("업로드할 파일\n"), 10, bold=True, color=MUTED)
    set_run_font(filename_paragraph.add_run(image_path.name), 10.5, bold=True, color=DARK)
    path_paragraph = right.add_paragraph()
    path_paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(path_paragraph.add_run(str(image_path)), 8.5, color=MUTED)

    footer_note = document.add_paragraph()
    footer_note.paragraph_format.space_before = Pt(8)
    footer_note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer_note.add_run(f"PPT {record['slideCount']}장 · 활동지 1장"), 8.5, color=MUTED)


def build_document() -> Path:
    register = json.loads(REGISTER_PATH.read_text(encoding="utf-8"))
    records = register["records"]
    if len(records) != 30:
        raise ValueError(f"등록 문구가 30개가 아닙니다: {len(records)}")
    intents = {record["teachingIntent"] for record in records}
    if len(intents) != 30:
        raise ValueError("등록용 수업 설계 의도가 중복되었습니다.")

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    normal = document.styles["Normal"]
    normal.font.name = DOCUMENT_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), DOCUMENT_FONT)
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    style_language = OxmlElement("w:lang")
    style_language.set(qn("w:val"), "ko-KR")
    style_language.set(qn("w:eastAsia"), "ko-KR")
    normal._element.rPr.append(style_language)
    normal.font.size = Pt(10.5)

    for index, record in enumerate(records):
        add_page(document, record, index, len(records))

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
